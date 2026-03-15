#!/usr/bin/env python3
"""
03_table1.py — Baseline characteristics table (converters vs non-converters)
"""

import pandas as pd
import numpy as np
import os, warnings
warnings.filterwarnings('ignore')
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SEED = 42
np.random.seed(SEED)

# NOTE: Update BASE to your local PPMI data directory
BASE = os.environ.get("PPMI_NOMOGRAM_DIR", ".")

# ─── Load ────────────────────────────────────────────────────────────────────
df = pd.read_csv(os.path.join(BASE, "data", "analytical_dataset.csv"))
df = df.dropna(subset=['time_years']).reset_index(drop=True)

# Recode variables
df['SAA_positive'] = (df['SAA_Status_Combined'].astype(str).str.lower().isin(
    ['positive', '1', '1.0', 'pos'])).astype(float)
df.loc[df['SAA_Status_Combined'].isna(), 'SAA_positive'] = np.nan

df['HVLT_total_learning'] = df[['HVLTRT1', 'HVLTRT2', 'HVLTRT3']].sum(axis=1, min_count=2)

df['APOE4_carrier'] = df['APOE'].astype(str).str.contains('E4', na=False).astype(float)
df.loc[df['APOE'].isna(), 'APOE4_carrier'] = np.nan

df['GBA_carrier'] = (~df['GBA'].isin(['0', 0]) & df['GBA'].notna()).astype(float)
df.loc[df['GBA'].isna(), 'GBA_carrier'] = np.nan

df['LRRK2_carrier'] = (~df['LRRK2'].isin(['0', 0]) & df['LRRK2'].notna()).astype(float)
df.loc[df['LRRK2'].isna(), 'LRRK2_carrier'] = np.nan

if 'SYSSUP' in df.columns:
    df['ortho_hypotension'] = ((df['SYSSUP'] - df['SYSSTND'] >= 20) |
                                (df['DIASUP'] - df['DIASTND'] >= 10)).astype(float)

if 'tremor_score' in df.columns:
    df['PIGD_dominant'] = (df['pigd_score'] > df['tremor_score']).astype(float)

print(f"N={len(df)}, Converters={df['event'].sum()}, Non-converters={(df['event']==0).sum()}")

# ─── Groups ──────────────────────────────────────────────────────────────────
converters = df[df['event'] == 1]
non_converters = df[df['event'] == 0]

# ─── Helper functions ────────────────────────────────────────────────────────
def fmt_continuous(data, decimals=1):
    """Mean ± SD"""
    d = data.dropna()
    if len(d) == 0:
        return "—"
    return f"{d.mean():.{decimals}f} ± {d.std():.{decimals}f}"

def fmt_continuous_median(data, decimals=1):
    """Median (IQR)"""
    d = data.dropna()
    if len(d) == 0:
        return "—"
    q25, q50, q75 = d.quantile([0.25, 0.5, 0.75])
    return f"{q50:.{decimals}f} ({q25:.{decimals}f}–{q75:.{decimals}f})"

def fmt_categorical(data, value=1):
    """N (%)"""
    d = data.dropna()
    n = (d == value).sum()
    pct = 100 * n / len(d) if len(d) > 0 else 0
    return f"{n} ({pct:.1f})"

def p_continuous(g1, g2):
    """t-test or Mann-Whitney"""
    d1, d2 = g1.dropna(), g2.dropna()
    if len(d1) < 3 or len(d2) < 3:
        return "—"
    # Shapiro test for normality (skip if N > 5000)
    try:
        stat, p = stats.mannwhitneyu(d1, d2, alternative='two-sided')
        return format_p(p)
    except:
        return "—"

def p_categorical(g1, g2, value=1):
    """Chi-square test"""
    d1, d2 = g1.dropna(), g2.dropna()
    table = np.array([
        [(d1 == value).sum(), (d1 != value).sum()],
        [(d2 == value).sum(), (d2 != value).sum()]
    ])
    if table.min() < 0 or table.sum() == 0:
        return "—"
    try:
        if table.min() < 5:
            _, p = stats.fisher_exact(table)
        else:
            _, p, _, _ = stats.chi2_contingency(table)
        return format_p(p)
    except:
        return "—"

def format_p(p):
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"

def missing_str(data):
    n_miss = data.isna().sum()
    return f"{n_miss}" if n_miss > 0 else ""

# ─── Build table rows ────────────────────────────────────────────────────────
rows = []

# Demographics
rows.append(("Demographics", "", "", "", "", ""))
rows.append(("  Age, years", fmt_continuous(df['AGE_AT_VISIT']),
             fmt_continuous(converters['AGE_AT_VISIT']),
             fmt_continuous(non_converters['AGE_AT_VISIT']),
             p_continuous(converters['AGE_AT_VISIT'], non_converters['AGE_AT_VISIT']),
             missing_str(df['AGE_AT_VISIT'])))

rows.append(("  Male sex, n (%)", fmt_categorical(df['SEX']),
             fmt_categorical(converters['SEX']),
             fmt_categorical(non_converters['SEX']),
             p_categorical(converters['SEX'], non_converters['SEX']),
             missing_str(df['SEX'])))

rows.append(("  Education, years", fmt_continuous(df['EDUCYRS']),
             fmt_continuous(converters['EDUCYRS']),
             fmt_continuous(non_converters['EDUCYRS']),
             p_continuous(converters['EDUCYRS'], non_converters['EDUCYRS']),
             missing_str(df['EDUCYRS'])))

rows.append(("  Follow-up, years", fmt_continuous_median(df['time_years']),
             fmt_continuous_median(converters['time_years']),
             fmt_continuous_median(non_converters['time_years']),
             p_continuous(converters['time_years'], non_converters['time_years']),
             ""))

# Genetics
rows.append(("Genetics", "", "", "", "", ""))
rows.append(("  APOE ε4 carrier, n (%)", fmt_categorical(df['APOE4_carrier']),
             fmt_categorical(converters['APOE4_carrier']),
             fmt_categorical(non_converters['APOE4_carrier']),
             p_categorical(converters['APOE4_carrier'], non_converters['APOE4_carrier']),
             missing_str(df['APOE4_carrier'])))

rows.append(("  GBA mutation, n (%)", fmt_categorical(df['GBA_carrier']),
             fmt_categorical(converters['GBA_carrier']),
             fmt_categorical(non_converters['GBA_carrier']),
             p_categorical(converters['GBA_carrier'], non_converters['GBA_carrier']),
             missing_str(df['GBA_carrier'])))

rows.append(("  LRRK2 mutation, n (%)", fmt_categorical(df['LRRK2_carrier']),
             fmt_categorical(converters['LRRK2_carrier']),
             fmt_categorical(non_converters['LRRK2_carrier']),
             p_categorical(converters['LRRK2_carrier'], non_converters['LRRK2_carrier']),
             missing_str(df['LRRK2_carrier'])))

# Motor
rows.append(("Motor assessments", "", "", "", "", ""))
rows.append(("  MDS-UPDRS Part III", fmt_continuous(df['NP3TOT_BEST']),
             fmt_continuous(converters['NP3TOT_BEST']),
             fmt_continuous(non_converters['NP3TOT_BEST']),
             p_continuous(converters['NP3TOT_BEST'], non_converters['NP3TOT_BEST']),
             missing_str(df['NP3TOT_BEST'])))

rows.append(("  MDS-UPDRS Part II", fmt_continuous(df['NP2PTOT']),
             fmt_continuous(converters['NP2PTOT']),
             fmt_continuous(non_converters['NP2PTOT']),
             p_continuous(converters['NP2PTOT'], non_converters['NP2PTOT']),
             missing_str(df['NP2PTOT'])))

rows.append(("  MDS-UPDRS Part I", fmt_continuous(df['NP1RTOT']),
             fmt_continuous(converters['NP1RTOT']),
             fmt_continuous(non_converters['NP1RTOT']),
             p_continuous(converters['NP1RTOT'], non_converters['NP1RTOT']),
             missing_str(df['NP1RTOT'])))

rows.append(("  H&Y stage", fmt_continuous_median(df['NHY_BEST']),
             fmt_continuous_median(converters['NHY_BEST']),
             fmt_continuous_median(non_converters['NHY_BEST']),
             p_continuous(converters['NHY_BEST'], non_converters['NHY_BEST']),
             missing_str(df['NHY_BEST'])))

rows.append(("  PIGD dominant, n (%)", fmt_categorical(df['PIGD_dominant']),
             fmt_categorical(converters['PIGD_dominant']),
             fmt_categorical(non_converters['PIGD_dominant']),
             p_categorical(converters['PIGD_dominant'], non_converters['PIGD_dominant']),
             missing_str(df['PIGD_dominant'])))

# Cognitive tests
rows.append(("Cognitive assessments", "", "", "", "", ""))
rows.append(("  MoCA", fmt_continuous(df['MCATOT']),
             fmt_continuous(converters['MCATOT']),
             fmt_continuous(non_converters['MCATOT']),
             p_continuous(converters['MCATOT'], non_converters['MCATOT']),
             missing_str(df['MCATOT'])))

rows.append(("  HVLT-R delayed recall", fmt_continuous(df['HVLTRDLY']),
             fmt_continuous(converters['HVLTRDLY']),
             fmt_continuous(non_converters['HVLTRDLY']),
             p_continuous(converters['HVLTRDLY'], non_converters['HVLTRDLY']),
             missing_str(df['HVLTRDLY'])))

rows.append(("  HVLT-R total learning", fmt_continuous(df['HVLT_total_learning']),
             fmt_continuous(converters['HVLT_total_learning']),
             fmt_continuous(non_converters['HVLT_total_learning']),
             p_continuous(converters['HVLT_total_learning'], non_converters['HVLT_total_learning']),
             missing_str(df['HVLT_total_learning'])))

rows.append(("  JLO total", fmt_continuous(df['JLO_TOTRAW']),
             fmt_continuous(converters['JLO_TOTRAW']),
             fmt_continuous(non_converters['JLO_TOTRAW']),
             p_continuous(converters['JLO_TOTRAW'], non_converters['JLO_TOTRAW']),
             missing_str(df['JLO_TOTRAW'])))

rows.append(("  Letter-Number Sequencing", fmt_continuous(df['LNS_TOTRAW']),
             fmt_continuous(converters['LNS_TOTRAW']),
             fmt_continuous(non_converters['LNS_TOTRAW']),
             p_continuous(converters['LNS_TOTRAW'], non_converters['LNS_TOTRAW']),
             missing_str(df['LNS_TOTRAW'])))

rows.append(("  Symbol Digit Modalities", fmt_continuous(df['SDMTOTAL']),
             fmt_continuous(converters['SDMTOTAL']),
             fmt_continuous(non_converters['SDMTOTAL']),
             p_continuous(converters['SDMTOTAL'], non_converters['SDMTOTAL']),
             missing_str(df['SDMTOTAL'])))

rows.append(("  Semantic fluency (animals)", fmt_continuous(df['DVT_SFTANIM']),
             fmt_continuous(converters['DVT_SFTANIM']),
             fmt_continuous(non_converters['DVT_SFTANIM']),
             p_continuous(converters['DVT_SFTANIM'], non_converters['DVT_SFTANIM']),
             missing_str(df['DVT_SFTANIM'])))

# NMS
rows.append(("Non-motor symptoms", "", "", "", "", ""))
rows.append(("  Cognitive complaints (1.1)", fmt_continuous_median(df['NP1COG']),
             fmt_continuous_median(converters['NP1COG']),
             fmt_continuous_median(non_converters['NP1COG']),
             p_continuous(converters['NP1COG'], non_converters['NP1COG']),
             missing_str(df['NP1COG'])))

rows.append(("  Hallucinations (1.2)", fmt_continuous_median(df['NP1HALL']),
             fmt_continuous_median(converters['NP1HALL']),
             fmt_continuous_median(non_converters['NP1HALL']),
             p_continuous(converters['NP1HALL'], non_converters['NP1HALL']),
             missing_str(df['NP1HALL'])))

rows.append(("  Depression (1.3)", fmt_continuous_median(df['NP1DPRS']),
             fmt_continuous_median(converters['NP1DPRS']),
             fmt_continuous_median(non_converters['NP1DPRS']),
             p_continuous(converters['NP1DPRS'], non_converters['NP1DPRS']),
             missing_str(df['NP1DPRS'])))

rows.append(("  Apathy (1.5)", fmt_continuous_median(df['NP1APAT']),
             fmt_continuous_median(converters['NP1APAT']),
             fmt_continuous_median(non_converters['NP1APAT']),
             p_continuous(converters['NP1APAT'], non_converters['NP1APAT']),
             missing_str(df['NP1APAT'])))

rows.append(("  RBD Screening Score", fmt_continuous(df['RBDSQ_TOTAL']),
             fmt_continuous(converters['RBDSQ_TOTAL']),
             fmt_continuous(non_converters['RBDSQ_TOTAL']),
             p_continuous(converters['RBDSQ_TOTAL'], non_converters['RBDSQ_TOTAL']),
             missing_str(df['RBDSQ_TOTAL'])))

rows.append(("  SCOPA-AUT total", fmt_continuous(df['SCAU_TOTAL']),
             fmt_continuous(converters['SCAU_TOTAL']),
             fmt_continuous(non_converters['SCAU_TOTAL']),
             p_continuous(converters['SCAU_TOTAL'], non_converters['SCAU_TOTAL']),
             missing_str(df['SCAU_TOTAL'])))

rows.append(("  Orthostatic hypotension, n (%)", fmt_categorical(df['ortho_hypotension']),
             fmt_categorical(converters['ortho_hypotension']),
             fmt_categorical(non_converters['ortho_hypotension']),
             p_categorical(converters['ortho_hypotension'], non_converters['ortho_hypotension']),
             missing_str(df['ortho_hypotension'])))

# Functional
rows.append(("Functional", "", "", "", "", ""))
rows.append(("  Modified Schwab-England ADL", fmt_continuous(df['MSEADLG']),
             fmt_continuous(converters['MSEADLG']),
             fmt_continuous(non_converters['MSEADLG']),
             p_continuous(converters['MSEADLG'], non_converters['MSEADLG']),
             missing_str(df['MSEADLG'])))

# Biomarkers
rows.append(("  α-synuclein SAA positive, n (%)", fmt_categorical(df['SAA_positive']),
             fmt_categorical(converters['SAA_positive']),
             fmt_categorical(non_converters['SAA_positive']),
             p_categorical(converters['SAA_positive'], non_converters['SAA_positive']),
             missing_str(df['SAA_positive'])))

# ─── Create Word document ───────────────────────────────────────────────────
doc = Document()

# Table style
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(9)

# Add table legend above
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(f'Table 1. Baseline characteristics of PD patients stratified by cognitive outcome (N = {len(df)})')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Arial'

# Create table
headers = ['Characteristic', f'All (N={len(df)})',
           f'Converters (N={len(converters)})',
           f'Non-converters (N={len(non_converters)})', 'p-value', 'Missing']
n_cols = len(headers)
table = doc.add_table(rows=1 + len(rows), cols=n_cols, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'

# Data rows
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.rows[i + 1].cells[j]
        cell.text = str(val)
        for paragraph in cell.paragraphs:
            if j == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Arial'
                # Bold section headers
                if val in ['Demographics', 'Genetics', 'Motor assessments',
                          'Cognitive assessments', 'Non-motor symptoms', 'Functional']:
                    run.bold = True

# Footnote
p = doc.add_paragraph()
run = p.add_run(
    'Values are mean ± SD, median (IQR), or n (%). p-values from Mann-Whitney U test '
    '(continuous) or chi-square/Fisher exact test (categorical). '
    'MDS-UPDRS, Movement Disorder Society–Unified Parkinson\'s Disease Rating Scale; '
    'H&Y, Hoehn and Yahr; PIGD, postural instability and gait difficulty; '
    'MoCA, Montreal Cognitive Assessment; HVLT-R, Hopkins Verbal Learning Test–Revised; '
    'JLO, Judgment of Line Orientation; SDM, Symbol Digit Modalities; '
    'RBD, rapid eye movement sleep behavior disorder; SCOPA-AUT, Scales for Outcomes in '
    'Parkinson\'s Disease–Autonomic; ADL, activities of daily living; '
    'SAA, seed amplification assay; LEDD, levodopa equivalent daily dose.'
)
run.font.size = Pt(8)
run.font.name = 'Arial'
run.italic = True

doc.save(os.path.join(BASE, "tables", "Table1.docx"))
print(f"Saved: tables/Table1.docx")

# Also save as CSV
table_df = pd.DataFrame(rows, columns=headers)
table_df.to_csv(os.path.join(BASE, "tables", "Table1.csv"), index=False)
print(f"Saved: tables/Table1.csv")
print("✓ Table 1 complete.")
