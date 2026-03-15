#!/usr/bin/env python3
"""
06_workflow_and_cover.py — Generate Workflow_Decisions.docx and cover letter
"""

import os, json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# NOTE: Update BASE to your local PPMI data directory
BASE = os.environ.get("PPMI_NOMOGRAM_DIR", ".")

with open(os.path.join(BASE, "logs", "model_results.json")) as f:
    cox_results = json.load(f)
with open(os.path.join(BASE, "logs", "cohort_info.json")) as f:
    cohort_info = json.load(f)

# ═══════════════════════════════════════════════════════════════════════
# WORKFLOW DECISIONS
# ═══════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

doc.add_heading('Workflow Decisions — PD Cognitive Impairment Nomogram', level=1)
doc.add_paragraph('Project: Predicting Cognitive Impairment in PD Using PPMI Data')
doc.add_paragraph('Date: March 2026')

doc.add_heading('1. Data Source and Population', level=2)
doc.add_paragraph(
    'Source: PPMI Master Merged dataset (PPMI_Master_Merged.csv)\n'
    'Total PPMI participants: 7,781\n'
    'PD cohort (COHORT=1): 1,970 patients\n\n'
    'Rationale: PPMI is the largest prospective, multicenter longitudinal study of de novo PD '
    'with standardized cognitive assessments, genetic testing, and biomarker collection.'
)

doc.add_heading('2. Inclusion/Exclusion Criteria', level=2)
doc.add_paragraph(
    f'Inclusion:\n'
    f'  - PD diagnosis (COHORT = 1)\n'
    f'  - At least one cognitive state assessment (COGSTATE available)\n'
    f'  - Normal cognition (COGSTATE = 1) at first assessment\n'
    f'  - ≥2 cognitive assessments (to detect events)\n'
    f'  - Available time-to-event data\n\n'
    f'Exclusion cascade (CONSORT):\n'
    f'  N = 7,781 → PD only → N = 1,970 (excluded 5,811 non-PD)\n'
    f'  N = 1,970 → With COGSTATE → N = 1,520 (excluded 450 no assessment)\n'
    f'  N = 1,520 → Normal at first → N = 1,328 (excluded 192 MCI/dementia)\n'
    f'  N = 1,328 → ≥2 assessments → N = 1,180 (excluded 148 single)\n'
    f'  N = 1,180 → With time data → N = 1,152 (excluded 28 missing)\n\n'
    'Rationale for each criterion:\n'
    '  - PD only: Nomogram is disease-specific\n'
    '  - Normal cognition at baseline: Incident prediction (not prevalence)\n'
    '  - ≥2 assessments: Minimum for time-to-event determination\n'
    '  - Excluding prevalent MCI/dementia avoids left censoring bias'
)

doc.add_heading('3. Outcome Definition', level=2)
doc.add_paragraph(
    'Primary outcome: Time to first cognitive impairment (COGSTATE ≥ 2)\n'
    '  - COGSTATE = 1: Normal cognition\n'
    '  - COGSTATE = 2: MCI (per MDS criteria)\n'
    '  - COGSTATE = 3: Dementia\n\n'
    'Time measurement: PPMI visit numbers mapped to years from baseline:\n'
    '  BL=0, V01=0.25, V02=0.5, V03=0.75, V04=1, V05=2, V06=3, V08=4, V10=5,\n'
    '  V12=6, V13=7, V14=8, V15=9, V16=10, V17=11, V18=12, V19=13, V20=14, V21=15\n\n'
    'Censoring: Patients without event censored at last COGSTATE assessment.\n'
    'Minimum time: 0.25 years (enforced for very short intervals).\n\n'
    'Rationale: COGSTATE incorporates MDS Level I and II criteria. Time-to-event '
    'framework is most appropriate for this longitudinal prediction question.'
)

doc.add_heading('4. Predictor Variable Selection', level=2)
doc.add_paragraph(
    '31 candidate predictors evaluated across 7 domains:\n\n'
    'Demographics (3): Age, sex, education years\n'
    'Genetics (3): APOE ε4, GBA, LRRK2 — coded as binary carrier status\n'
    '  - APOE: from genotype string (e.g., "E3/E4" → carrier)\n'
    '  - GBA: from variant name ("0" = wildtype, any other = carrier)\n'
    '  - LRRK2: same as GBA\n'
    'Motor (6): UPDRS-III best, H&Y best, UPDRS-II, UPDRS-I, tremor, PIGD\n'
    'Cognitive (7): MoCA, HVLT-R delayed, HVLT-R total learning (trials 1-3 sum),\n'
    '  JLO, LNS, SDMT, semantic fluency\n'
    'Neuropsychiatric (5): Cognitive complaints, hallucinations, depression, anxiety, apathy\n'
    'Non-motor (4): RBD score, SCOPA-AUT, orthostatic hypotension, PIGD dominant\n'
    'Functional (3): Schwab-England ADL, SAA positive, LEDD\n\n'
    'Rationale for variable derivations:\n'
    '  - HVLT total learning (sum of trials 1-3) used instead of individual trials\n'
    '    to reduce multicollinearity\n'
    '  - DVS_LNS (scaled score) dropped in favor of LNS raw score\n'
    '  - Orthostatic hypotension derived from BP measures (≥20 systolic or ≥10 diastolic drop)\n'
    '  - PIGD dominant from comparison of PIGD vs tremor composite scores\n\n'
    'Missing data: Median imputation. All primary candidates had ≥70% availability.\n'
    'Variables with <30% availability (DATScan, CSF NfL/GFAP) excluded from primary model.'
)

doc.add_heading('5. Analytical Decisions', level=2)
doc.add_paragraph(
    'Variable selection: LASSO-penalized Cox regression\n'
    '  - Elastic net mixing parameter: 0.9 (predominant L1)\n'
    '  - Alpha selection: 10-fold stratified CV, maximizing C-index\n'
    f'  - Optimal alpha: {cox_results["best_lasso_alpha"]:.4f}\n'
    f'  - Variables selected: {cox_results["n_predictors_selected"]} of {cox_results["n_predictors_initial"]}\n\n'
    'Final model: Cox PH with small ridge penalty (0.01) for stability\n\n'
    'Validation strategy:\n'
    '  1. 10-fold stratified CV: C-index\n'
    '  2. Bootstrap (200 iterations): Optimism correction\n'
    '  3. Calibration: Quintile-based at 3, 5, 8 years\n'
    '  4. Time-dependent AUC: Annual from 1-12 years\n\n'
    'Sensitivity analyses:\n'
    '  1. Gradient boosting survival analysis\n'
    '  2. Random survival forest\n'
    '  3. CSF/plasma biomarker-enhanced model (subset analysis)\n\n'
    'Random seed: 42 (used throughout all stochastic procedures)\n\n'
    'Rationale:\n'
    '  - LASSO preferred for automatic variable selection with built-in regularization\n'
    '  - Cox PH preferred over ML for nomogram interpretability\n'
    '  - Multiple validation approaches recommended by TRIPOD guidelines\n'
    '  - Biomarker analysis as sensitivity due to high missingness (>50%)'
)

doc.add_heading('6. Key Results Summary', level=2)
doc.add_paragraph(
    f'Final N: {cox_results["n_total"]}\n'
    f'Events: {cox_results["n_events"]} ({100*cox_results["n_events"]/cox_results["n_total"]:.1f}%)\n'
    f'Median follow-up: 4.0 years\n'
    f'Apparent C-index: {cox_results["apparent_c_index"]:.3f}\n'
    f'CV C-index: {cox_results["cv_c_index_mean"]:.3f} ± {cox_results["cv_c_index_std"]:.3f}\n'
    f'Optimism-corrected C-index: {cox_results["optimism_corrected_c_index"]:.3f}\n'
    f'Mean time-dependent AUC: 0.866\n'
    f'Log-rank (low vs high risk): p < 0.001'
)

doc.save(os.path.join(BASE, "Workflow_Decisions.docx"))
print("Saved: Workflow_Decisions.docx")

# ═══════════════════════════════════════════════════════════════════════
# COVER LETTER
# ═══════════════════════════════════════════════════════════════════════
cl = Document()
style = cl.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

cl.add_paragraph('March 15, 2026\n')

cl.add_paragraph('Dear Editor,\n')

cl.add_paragraph(
    'We are pleased to submit our manuscript entitled "Development and Validation of a Clinical '
    'Nomogram for Predicting Cognitive Impairment in Parkinson\'s Disease: A Longitudinal Cohort '
    'Study from the PPMI" for consideration for publication in your journal.'
)

cl.add_paragraph(
    'Cognitive impairment is the most impactful non-motor complication of Parkinson\'s disease, '
    'affecting up to 80% of patients over their disease course. Despite the identification of '
    'multiple risk factors, no validated, clinically practical tool exists for predicting which '
    'patients will develop cognitive impairment and when. Our study addresses this critical gap '
    'by developing a nomogram—a user-friendly graphical prediction tool—using 1,152 PD patients '
    'from the landmark PPMI longitudinal cohort with up to 15 years of follow-up.'
)

cl.add_paragraph(
    'Our nomogram integrates readily available clinical, cognitive, and genetic predictors '
    '(including GBA and LRRK2 mutation status) and achieves a time-dependent AUC consistently '
    'exceeding 0.80 at all evaluated timepoints (mean AUC = 0.87). The model effectively '
    'stratifies patients into distinct risk trajectories (log-rank p < 0.001) and performs '
    'comparably to machine learning approaches while maintaining full interpretability. '
    'Notably, CSF and plasma biomarkers provided minimal incremental value over clinical '
    'predictors, demonstrating that accurate prediction is achievable without invasive '
    'biomarker testing.'
)

cl.add_paragraph(
    'We believe this work is of particular relevance to your readership because: '
    '(1) it provides the first comprehensive nomogram for cognitive impairment in PD from a '
    'large multicenter cohort; (2) it highlights the dominant role of GBA mutations in cognitive '
    'risk, with important implications for genetic counseling and emerging GBA-targeted therapies; '
    'and (3) the nomogram offers immediate clinical applicability for risk stratification and '
    'clinical trial enrichment.'
)

cl.add_paragraph(
    'This manuscript has not been published previously, is not under consideration elsewhere, '
    'and all authors have approved the manuscript. There are no conflicts of interest to declare.\n'
)

cl.add_paragraph(
    'Word count: approximately 4,500 words\n'
    'Figures: 7 main figures\n'
    'Tables: 2 tables\n'
    'Supplementary figures: 1\n'
)

cl.add_paragraph('Sincerely,\n')

cl.add_paragraph(
    'Ahmed Negida, MD, PhD\n'
    'Department of Neurology\n'
    'Virginia Commonwealth University\n'
    '417 N 11th St, Richmond, VA 23298\n'
    'Email: ahmed.negida@vcuhealth.org\n'
    'Phone: +1 347-251-3472\n'
    'ORCID: 0000-0001-5363-6369'
)

cl.save(os.path.join(BASE, "cover_letter", "cover_letter.docx"))
print("Saved: cover_letter/cover_letter.docx")
print("✓ All documents generated.")
